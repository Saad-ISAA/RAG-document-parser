# parsers/document_parser.py
"""
Enhanced document parser for DOCX, DOC, ODT, RTF and other word processing documents.
Now includes improved DOC support with conversion fallback and better Arabic text handling.
"""

import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from datetime import datetime
import io
import tempfile

# Import document libraries with fallbacks
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    DocxDocument = None

try:
    from odf import text, teletype
    from odf.opendocument import load as odf_load
    HAS_ODFPY = True
except ImportError:
    HAS_ODFPY = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False

try:
    import olefile
    HAS_OLEFILE = True
except ImportError:
    HAS_OLEFILE = False

# Enhanced DOC parsing
try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

# For Arabic text processing
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    HAS_ARABIC_SUPPORT = True
except ImportError:
    HAS_ARABIC_SUPPORT = False

from models.parse_result import ParseResult, TableData, ImageInfo, DocumentMetadata, FileInfo
from utils.config import ParserConfig

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Enhanced parser for various document formats including DOCX, DOC, ODT, RTF.
    Now with improved DOC support and Arabic text handling.
    """
    
    def __init__(self, config: ParserConfig):
        """Initialize document parser with configuration."""
        self.config = config
        self.doc_config = config.document
        
        # Check available libraries
        self.available_libs = {}
        if HAS_PYTHON_DOCX:
            self.available_libs['docx'] = True
        if HAS_ODFPY:
            self.available_libs['odt'] = True
        if HAS_STRIPRTF:
            self.available_libs['rtf'] = True
        if HAS_OLEFILE:
            self.available_libs['doc'] = True
        if HAS_MAMMOTH:
            self.available_libs['mammoth'] = True
        
        logger.info(f"Document parser initialized with support for: {', '.join(self.available_libs.keys())}")
        
        if HAS_ARABIC_SUPPORT:
            logger.info("Arabic text processing support available")
        else:
            logger.info("Arabic text processing not available (install arabic-reshaper and python-bidi)")
    
    def parse(self, file_path: Path, file_info: FileInfo) -> ParseResult:
        """
        Parse a document file based on its type.
        
        Args:
            file_path: Path to the document file
            file_info: File information from detector
            
        Returns:
            ParseResult with extracted content
        """
        result = ParseResult(
            file_path=str(file_path),
            file_type=file_info.mime_type
        )
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.docx':
                return self._parse_docx(file_path, result)
            elif extension == '.doc':
                return self._parse_doc_enhanced(file_path, result)
            elif extension in ['.odt']:
                return self._parse_odt(file_path, result)
            elif extension == '.rtf':
                return self._parse_rtf(file_path, result)
            else:
                result.error = f"Unsupported document format: {extension}"
                return result
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            result.error = str(e)
            return result
    
    def _parse_doc_enhanced(self, file_path: Path, result: ParseResult) -> ParseResult:
        """Enhanced DOC parsing with multiple methods and conversion fallback."""
        
        # Method 1: Try Mammoth (best for DOC files)
        if HAS_MAMMOTH:
            mammoth_result = self._parse_doc_with_mammoth(file_path, result)
            if mammoth_result.success:
                return mammoth_result
        
        # Method 2: Try conversion to DOCX, then parse
        try:
            from utils.document_converter import DocumentConverter
            converter = DocumentConverter(self.config)
            
            # Convert DOC to PDF first, then to DOCX if possible
            if converter.can_convert(file_path, 'pdf'):
                logger.info(f"Attempting conversion of {file_path.name} to PDF for better parsing")
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = Path(temp_dir)
                    
                    # Convert to PDF
                    pdf_file = converter.convert_to_pdf(file_path, temp_path)
                    if pdf_file and pdf_file.exists():
                        # Parse the PDF
                        from parsers.pdf_parser import PDFParser
                        pdf_parser = PDFParser(self.config)
                        
                        # Create file info for PDF
                        from models.parse_result import FileInfo
                        pdf_info = FileInfo(
                            mime_type='application/pdf',
                            category='pdf',
                            extension='.pdf'
                        )
                        
                        pdf_result = pdf_parser.parse(pdf_file, pdf_info)
                        if pdf_result.success:
                            pdf_result.parser_used = 'doc-conversion-pdf'
                            pdf_result.file_path = str(file_path)  # Keep original path
                            return pdf_result
        
        except ImportError:
            logger.debug("Document converter not available for DOC parsing")
        except Exception as e:
            logger.warning(f"DOC conversion failed: {str(e)}")
        
        # Method 3: Basic binary extraction (fallback)
        return self._parse_doc_basic(file_path, result)
    
    def _parse_doc_with_mammoth(self, file_path: Path, result: ParseResult) -> ParseResult:
        """Parse DOC file using Mammoth library."""
        try:
            with open(file_path, 'rb') as docx_file:
                doc_result = mammoth.extract_text(docx_file)
                
                result.content = doc_result.value
                
                # Process Arabic text if needed
                result.content = self._process_arabic_text(result.content)
                
                # Handle warnings
                if doc_result.messages:
                    warnings = [msg.message for msg in doc_result.messages]
                    logger.debug(f"Mammoth parsing warnings: {warnings}")
                
                # Basic metadata
                result.metadata = DocumentMetadata()
                result.metadata.word_count = len(result.content.split())
                result.metadata.page_count = max(1, len(result.content) // 2000)  # Estimate
                
                result.success = True
                result.parser_used = 'mammoth'
                
        except Exception as e:
            result.error = f"Mammoth DOC parsing failed: {str(e)}"
            logger.debug(f"Mammoth parsing error: {str(e)}")
        
        return result
    
    def _parse_doc_basic(self, file_path: Path, result: ParseResult) -> ParseResult:
        """Basic DOC parsing (fallback method)."""
        try:
            # This is a very basic approach - improved from original
            with open(file_path, 'rb') as f:
                content = f.read()
                
                # Try to decode with different encodings
                text_content = ""
                encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
                
                for encoding in encodings:
                    try:
                        decoded = content.decode(encoding, errors='ignore')
                        # Filter out binary content - improved filtering
                        clean_chars = []
                        for char in decoded:
                            if char.isprintable() or char in ['\n', '\r', '\t']:
                                clean_chars.append(char)
                        
                        potential_text = ''.join(clean_chars)
                        
                        # Only use if we got substantial readable text
                        words = potential_text.split()
                        readable_words = [w for w in words if len(w) > 1 and any(c.isalpha() for c in w)]
                        
                        if len(readable_words) > 50:  # Threshold for meaningful content
                            text_content = potential_text
                            break
                            
                    except Exception:
                        continue
                
                if text_content:
                    # Clean up the extracted text
                    lines = text_content.split('\n')
                    clean_lines = []
                    
                    for line in lines:
                        line = line.strip()
                        if line and len(line) > 2:  # Skip very short lines
                            # Remove lines that are mostly non-alphabetic
                            alpha_ratio = sum(c.isalpha() or c.isspace() for c in line) / len(line)
                            if alpha_ratio > 0.5:  # At least 50% alphabetic or space
                                clean_lines.append(line)
                    
                    result.content = '\n'.join(clean_lines)
                    result.content = self._process_arabic_text(result.content)
                    
                    if len(result.content.strip()) > 100:  # Minimum meaningful content
                        result.success = True
                        result.parser_used = 'basic-binary-enhanced'
                    else:
                        result.error = "DOC file contains insufficient readable text"
                else:
                    result.error = "Unable to extract readable text from DOC file"
        
        except Exception as e:
            result.error = f"Basic DOC parsing failed: {str(e)}"
        
        return result
    
    def _parse_docx(self, file_path: Path, result: ParseResult) -> ParseResult:
        """Parse DOCX file using python-docx with Arabic support."""
        if not HAS_PYTHON_DOCX:
            result.error = "python-docx library not available"
            return result
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            result.metadata = self._extract_docx_metadata(doc)
            
            # Extract text content
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Process Arabic text if needed
                    processed_text = self._process_arabic_text(para.text)
                    paragraphs.append(processed_text)
            
            result.content = '\n'.join(paragraphs)
            
            # Extract tables if enabled
            if self.doc_config.extract_tables:
                result.tables = self._extract_docx_tables(doc)
            
            # Extract images if enabled
            if self.doc_config.extract_images:
                result.images = self._extract_docx_images_enhanced(doc)
            
            # Extract headers and footers if enabled
            if self.doc_config.include_headers_footers:
                header_footer_text = self._extract_docx_headers_footers(doc)
                if header_footer_text:
                    result.content = header_footer_text + '\n\n' + result.content
            
            result.success = True
            result.parser_used = 'python-docx'
            
        except Exception as e:
            result.error = f"DOCX parsing failed: {str(e)}"
            logger.error(f"DOCX parsing error: {str(e)}")
        
        return result
    
    def _extract_docx_images_enhanced(self, doc) -> List[ImageInfo]:
        """Enhanced image extraction from DOCX with OCR support."""
        images = []
        
        try:
            # Get document relationships
            rels = doc.part.rels
            
            image_index = 0
            for rel in rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_info = ImageInfo(
                            image_index=image_index,
                            format=image_part.content_type.split('/')[-1],
                            size_bytes=len(image_part.blob)
                        )
                        
                        # Run OCR on image if enabled
                        if self.config.enable_ocr:
                            ocr_text = self._extract_text_from_image_blob(image_part.blob)
                            if ocr_text:
                                image_info.extracted_text = ocr_text
                                logger.debug(f"OCR extracted {len(ocr_text)} chars from DOCX image {image_index}")
                        
                        images.append(image_info)
                        image_index += 1
                        
                    except Exception as e:
                        logger.warning(f"Failed to process DOCX image {image_index}: {str(e)}")
        
        except Exception as e:
            logger.warning(f"Failed to extract DOCX images: {str(e)}")
        
        return images
    
    def _extract_text_from_image_blob(self, image_blob: bytes) -> Optional[str]:
        """Extract text from image blob using OCR."""
        if not self.config.enable_ocr:
            return None
        
        try:
            from PIL import Image
            import io
            
            # Convert blob to PIL Image
            pil_image = Image.open(io.BytesIO(image_blob))
            
            # Use EasyOCR if available
            try:
                import easyocr
                import numpy as np
                
                # Initialize with Arabic and English
                languages = list(set(self.config.ocr.languages + ['en', 'ar']))
                reader = easyocr.Reader(languages, verbose=False)
                
                # Convert to numpy array
                if pil_image.mode not in ['RGB', 'L']:
                    pil_image = pil_image.convert('RGB')
                image_array = np.array(pil_image)
                
                # Run OCR
                results = reader.readtext(image_array)
                
                # Extract text with confidence filtering
                text_parts = []
                for bbox, text, confidence in results:
                    if confidence >= self.config.ocr.confidence_threshold:
                        text_parts.append(text.strip())
                
                if text_parts:
                    return '\n'.join(text_parts)
                    
            except ImportError:
                # Fallback to Tesseract
                try:
                    import pytesseract
                    
                    # Configure for Arabic/English
                    lang = '+'.join(self.config.ocr.languages)
                    text = pytesseract.image_to_string(pil_image, lang=lang)
                    return text.strip() if text.strip() else None
                    
                except ImportError:
                    logger.debug("No OCR engines available for image text extraction")
        
        except Exception as e:
            logger.debug(f"Image OCR failed: {str(e)}")
        
        return None
    
    def _process_arabic_text(self, text: str) -> str:
        """Process Arabic text for proper display and formatting."""
        if not text or not HAS_ARABIC_SUPPORT:
            return text
        
        try:
            # Check if text contains Arabic characters
            has_arabic = any('\u0600' <= char <= '\u06FF' for char in text)
            
            if has_arabic:
                # Reshape Arabic text for proper display
                reshaped_text = arabic_reshaper.reshape(text)
                # Apply bidirectional algorithm
                display_text = get_display(reshaped_text)
                return display_text
            
        except Exception as e:
            logger.debug(f"Arabic text processing failed: {str(e)}")
        
        return text
    
    def _parse_odt(self, file_path: Path, result: ParseResult) -> ParseResult:
        """Parse ODT file using odfpy with Arabic support."""
        if not HAS_ODFPY:
            result.error = "odfpy library not available for ODT parsing"
            return result
        
        try:
            doc = odf_load(file_path)
            
            # Extract text content
            paragraphs = []
            for paragraph in doc.getElementsByType(text.P):
                para_text = teletype.extractText(paragraph)
                if para_text.strip():
                    # Process Arabic text
                    processed_text = self._process_arabic_text(para_text)
                    paragraphs.append(processed_text)
            
            result.content = '\n'.join(paragraphs)
            
            # Extract basic metadata
            result.metadata = self._extract_odt_metadata(doc)
            
            # TODO: Add table and image extraction for ODT
            
            result.success = True
            result.parser_used = 'odfpy'
            
        except Exception as e:
            result.error = f"ODT parsing failed: {str(e)}"
            logger.error(f"ODT parsing error: {str(e)}")
        
        return result
    
    def _parse_rtf(self, file_path: Path, result: ParseResult) -> ParseResult:
        """Parse RTF file using striprtf with Arabic support."""
        if not HAS_STRIPRTF:
            result.error = "striprtf library not available for RTF parsing"
            return result
        
        try:
            # Try different encodings for RTF files
            encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
            rtf_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        rtf_content = f.read()
                    break
                except Exception:
                    continue
            
            if not rtf_content:
                result.error = "Failed to read RTF file with any encoding"
                return result
            
            # Convert RTF to plain text
            plain_text = rtf_to_text(rtf_content)
            
            # Process Arabic text
            result.content = self._process_arabic_text(plain_text)
            
            # Basic metadata
            result.metadata = DocumentMetadata()
            result.metadata.page_count = 1  # RTF doesn't have explicit pages
            result.metadata.word_count = len(result.content.split())
            
            result.success = True
            result.parser_used = 'striprtf'
            
        except Exception as e:
            result.error = f"RTF parsing failed: {str(e)}"
            logger.error(f"RTF parsing error: {str(e)}")
        
        return result
    
    def _extract_docx_metadata(self, doc) -> DocumentMetadata:
        """Extract metadata from DOCX document."""
        metadata = DocumentMetadata()
        
        try:
            core_props = doc.core_properties
            
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.subject = core_props.subject
            metadata.creator = core_props.author
            metadata.creation_date = core_props.created
            metadata.modification_date = core_props.modified
            
            # Count words (approximate)
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())
            metadata.word_count = word_count
            
            # Page count is harder to determine without rendering
            # This is an approximation
            char_count = sum(len(para.text) for para in doc.paragraphs)
            metadata.page_count = max(1, char_count // 2000)  # Rough estimate
            
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {str(e)}")
        
        return metadata
    
    def _extract_docx_tables(self, doc) -> List[TableData]:
        """Extract tables from DOCX document."""
        tables = []
        
        try:
            for table_index, table in enumerate(doc.tables):
                if not table.rows:
                    continue
                
                # Extract headers (first row)
                headers = []
                if table.rows:
                    header_row = table.rows[0]
                    headers = [self._process_arabic_text(cell.text.strip()) for cell in header_row.cells]
                
                # Extract data rows
                rows = []
                for row in table.rows[1:]:  # Skip header row
                    row_data = [self._process_arabic_text(cell.text.strip()) for cell in row.cells]
                    rows.append(row_data)
                
                if headers and rows:
                    tables.append(TableData(
                        headers=headers,
                        rows=rows,
                        table_index=table_index
                    ))
        
        except Exception as e:
            logger.warning(f"Failed to extract DOCX tables: {str(e)}")
        
        return tables
    
    def _extract_docx_headers_footers(self, doc) -> str:
        """Extract headers and footers from DOCX document."""
        header_footer_text = []
        
        try:
            # Extract headers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            processed_text = self._process_arabic_text(para.text)
                            header_footer_text.append(f"[HEADER] {processed_text}")
                
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            processed_text = self._process_arabic_text(para.text)
                            header_footer_text.append(f"[FOOTER] {processed_text}")
        
        except Exception as e:
            logger.warning(f"Failed to extract headers/footers: {str(e)}")
        
        return '\n'.join(header_footer_text)
    
    def _extract_odt_metadata(self, doc) -> DocumentMetadata:
        """Extract metadata from ODT document."""
        metadata = DocumentMetadata()
        
        try:
            # ODT metadata extraction is more complex
            # This is a basic implementation
            meta = doc.getElementsByType(text.P)
            if meta:
                char_count = sum(len(teletype.extractText(p)) for p in meta)
                metadata.page_count = max(1, char_count // 2000)
                metadata.word_count = char_count // 5  # Rough word estimate
        
        except Exception as e:
            logger.warning(f"Failed to extract ODT metadata: {str(e)}")
        
        return metadata
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        formats = []
        if HAS_PYTHON_DOCX:
            formats.append('docx')
        if HAS_ODFPY:
            formats.append('odt')
        if HAS_STRIPRTF:
            formats.append('rtf')
        formats.append('doc')  # Enhanced support
        
        return formats
    
    def get_library_info(self) -> Dict[str, Any]:
        """Get information about available document libraries."""
        return {
            'available_libraries': self.available_libs,
            'arabic_support': HAS_ARABIC_SUPPORT,
            'capabilities': {
                'docx': {
                    'text_extraction': HAS_PYTHON_DOCX,
                    'table_extraction': HAS_PYTHON_DOCX,
                    'image_extraction': HAS_PYTHON_DOCX,
                    'metadata_extraction': HAS_PYTHON_DOCX,
                    'arabic_processing': HAS_ARABIC_SUPPORT
                },
                'doc': {
                    'text_extraction': True,  # Enhanced with multiple methods
                    'table_extraction': HAS_MAMMOTH,
                    'image_extraction': False,
                    'metadata_extraction': HAS_MAMMOTH,
                    'conversion_fallback': True
                },
                'odt': {
                    'text_extraction': HAS_ODFPY,
                    'table_extraction': False,  # TODO: Implement
                    'image_extraction': False,  # TODO: Implement
                    'metadata_extraction': HAS_ODFPY,
                    'arabic_processing': HAS_ARABIC_SUPPORT
                },
                'rtf': {
                    'text_extraction': HAS_STRIPRTF,
                    'table_extraction': False,
                    'image_extraction': False,
                    'metadata_extraction': False,
                    'arabic_processing': HAS_ARABIC_SUPPORT
                }
            }
        }